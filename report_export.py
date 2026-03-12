"""
稽核報告匯出模組

支援將稽核結果匯出為 Word (.docx) 和 PDF 格式。

用法:
    python report_export.py --input audit_reports/ --format docx
    python report_export.py --input audit_reports/batch_summary_xxx.json --format pdf
    python report_export.py --input audit_reports/ --format both
"""

import argparse
import json
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Union

sys.path.insert(0, str(Path(__file__).parent))


# ==========================================
# 資料結構
# ==========================================

@dataclass
class ReportData:
    """正規化的報告資料"""

    title: str
    generated_at: str
    rules_audited: int
    summary_stats: Dict[str, int] = field(default_factory=dict)
    severity_stats: Dict[str, int] = field(default_factory=dict)
    results: List[Dict] = field(default_factory=list)


# ==========================================
# 資料載入
# ==========================================

def load_audit_results(source) -> ReportData:
    """
    從多種來源載入稽核結果

    Args:
        source: 可以是:
            - Path 指向 audit_reports/ 目錄（讀取所有 JSON）
            - Path 指向 batch_summary JSON 檔案
            - List[Path] JSON 檔案列表
            - BatchAuditSummary 物件

    Returns:
        ReportData
    """
    # BatchAuditSummary 物件
    if hasattr(source, "results") and hasattr(source, "total_rules"):
        return _load_from_batch_summary(source)

    source = Path(source) if not isinstance(source, Path) else source

    # 目錄：讀取所有稽核 JSON
    if source.is_dir():
        return _load_from_directory(source)

    # 單一 JSON 檔案
    if source.is_file() and source.suffix == ".json":
        return _load_from_json_file(source)

    raise ValueError(f"無法辨識的來源: {source}")


def _load_from_directory(directory: Path) -> ReportData:
    """從目錄載入所有稽核報告 JSON"""
    json_files = sorted(directory.glob("稽核_RULE_*.json"))

    if not json_files:
        raise FileNotFoundError(f"在 {directory} 中未找到稽核報告 JSON")

    results = []
    for f in json_files:
        with open(f, "r", encoding="utf-8") as fp:
            data = json.load(fp)
            results.append(data)

    return _build_report_data(results)


def _load_from_json_file(filepath: Path) -> ReportData:
    """從 batch_summary 或單一報告 JSON 載入"""
    with open(filepath, "r", encoding="utf-8") as f:
        data = json.load(f)

    # batch summary 格式
    if "batch_id" in data and "results" in data:
        return _build_report_data_from_summary(data)

    # 單一報告格式
    return _build_report_data([data])


def _load_from_batch_summary(summary) -> ReportData:
    """從 BatchAuditSummary 物件載入"""
    results = []
    for r in summary.results:
        results.append({
            "rule": r.rule,
            "query_generation": r.query_info,
            "local_search_result": r.rag_result,
            "compliance_analysis": r.compliance,
        })
    return _build_report_data(results)


def _build_report_data(results: List[Dict]) -> ReportData:
    """從報告列表建立 ReportData"""
    stats = {"RISK_DETECTED": 0, "COMPLIANT": 0, "UNCERTAIN": 0, "NOT_APPLICABLE": 0, "ERROR": 0}
    severity = {"HIGH": 0, "MEDIUM": 0, "LOW": 0}

    for r in results:
        compliance = r.get("compliance_analysis", {})
        status = compliance.get("status", "ERROR")
        sev = compliance.get("severity", "LOW")
        stats[status] = stats.get(status, 0) + 1
        severity[sev] = severity.get(sev, 0) + 1

    return ReportData(
        title="契約稽核報告",
        generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        rules_audited=len(results),
        summary_stats=stats,
        severity_stats=severity,
        results=results,
    )


def _build_report_data_from_summary(data: Dict) -> ReportData:
    """從 batch summary JSON 建立 ReportData"""
    statistics = data.get("statistics", {})
    return ReportData(
        title="契約稽核報告",
        generated_at=data.get("end_time", datetime.now().isoformat()),
        rules_audited=data.get("total_rules", 0),
        summary_stats={
            "RISK_DETECTED": statistics.get("risk_detected", 0),
            "COMPLIANT": statistics.get("compliant", 0),
            "UNCERTAIN": statistics.get("uncertain", 0),
            "NOT_APPLICABLE": statistics.get("not_applicable", 0),
        },
        severity_stats={
            "HIGH": statistics.get("high_severity", 0),
            "MEDIUM": statistics.get("medium_severity", 0),
        },
        results=[],  # batch summary 只有摘要，無完整結果
    )


# ==========================================
# Word 匯出
# ==========================================

def export_to_docx(data: ReportData, output_path: Path) -> Path:
    """
    匯出稽核結果為 Word 文件

    Args:
        data: 正規化的報告資料
        output_path: 輸出檔案路徑

    Returns:
        輸出檔案路徑
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # 設定預設字型（CJK 支援）
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Noto Sans CJK TC"
    font.size = Pt(10)

    # ---- 封面 ----
    title = doc.add_heading(data.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"產出時間: {data.generated_at}\n").font.size = Pt(11)
    meta.add_run(f"稽核規則數: {data.rules_audited}\n").font.size = Pt(11)

    doc.add_page_break()

    # ---- 統計摘要 ----
    doc.add_heading("一、稽核統計摘要", level=1)

    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = "Light Grid Accent 1"
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = stats_table.rows[0].cells
    hdr[0].text = "項目"
    hdr[1].text = "數量"

    for label, count in [
        ("RISK_DETECTED (風險偵測)", data.summary_stats.get("RISK_DETECTED", 0)),
        ("COMPLIANT (合規)", data.summary_stats.get("COMPLIANT", 0)),
        ("UNCERTAIN (不確定)", data.summary_stats.get("UNCERTAIN", 0)),
        ("NOT_APPLICABLE (不適用)", data.summary_stats.get("NOT_APPLICABLE", 0)),
        ("HIGH 風險", data.severity_stats.get("HIGH", 0)),
        ("MEDIUM 風險", data.severity_stats.get("MEDIUM", 0)),
        ("LOW 風險", data.severity_stats.get("LOW", 0)),
    ]:
        row = stats_table.add_row().cells
        row[0].text = label
        row[1].text = str(count)

    doc.add_paragraph()

    # ---- 摘要表格 ----
    if data.results:
        doc.add_heading("二、逐規則摘要", level=1)

        summary_table = doc.add_table(rows=1, cols=5)
        summary_table.style = "Light Grid Accent 1"
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = summary_table.rows[0].cells
        hdr[0].text = "規則"
        hdr[1].text = "檢核項目"
        hdr[2].text = "判定結果"
        hdr[3].text = "風險等級"
        hdr[4].text = "條款參照"

        for result in data.results:
            rule = result.get("rule", {})
            compliance = result.get("compliance_analysis", {})
            row = summary_table.add_row().cells
            row[0].text = rule.get("id", "")
            row[1].text = rule.get("category", "")[:20]
            row[2].text = compliance.get("status", "")
            row[3].text = compliance.get("severity", "")
            row[4].text = compliance.get("clause_reference", "")

            # 狀態色彩標記
            status = compliance.get("status", "")
            status_cell = row[2]
            for paragraph in status_cell.paragraphs:
                for run in paragraph.runs:
                    if status == "RISK_DETECTED":
                        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                        run.bold = True
                    elif status == "COMPLIANT":
                        run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

        doc.add_page_break()

        # ---- 逐規則詳細 ----
        doc.add_heading("三、逐規則詳細報告", level=1)

        for idx, result in enumerate(data.results, 1):
            rule = result.get("rule", {})
            query = result.get("query_generation", {})
            search = result.get("local_search_result", {})
            compliance = result.get("compliance_analysis", {})

            doc.add_heading(
                f"{rule.get('id', '')} - {rule.get('category', '')}",
                level=2,
            )

            # 規則資訊
            doc.add_heading("規則資訊", level=3)
            doc.add_paragraph(f"錯誤樣態: {rule.get('risk_pattern', '')}")
            doc.add_paragraph(f"應採行動: {rule.get('action', '')}")
            doc.add_paragraph(f"標準釋疑: {rule.get('explanation', '')}")

            # 查詢資訊
            if query:
                doc.add_heading("查詢資訊", level=3)
                doc.add_paragraph(f"查詢語句: {query.get('graph_query', '')}")
                key_terms = query.get("key_terms", [])
                if key_terms:
                    doc.add_paragraph(f"關鍵詞: {', '.join(key_terms)}")

            # 檢索結果
            if search:
                doc.add_heading("檢索結果", level=3)
                answer = search.get("answer", "")
                # 截斷過長的回答
                if len(answer) > 2000:
                    answer = answer[:2000] + "\n...(內容已截斷)"
                doc.add_paragraph(answer)

            # 合規分析
            doc.add_heading("合規分析", level=3)
            doc.add_paragraph(f"判定結果: {compliance.get('status', '')}")
            doc.add_paragraph(f"風險等級: {compliance.get('severity', '')}")
            doc.add_paragraph(f"判斷理由: {compliance.get('reason', '')}")
            doc.add_paragraph(f"契約證據: {compliance.get('evidence', '')}")
            doc.add_paragraph(f"條款參照: {compliance.get('clause_reference', '')}")
            doc.add_paragraph(f"建議行動: {compliance.get('recommendation', '')}")

            # 分隔（非最後一項）
            if idx < len(data.results):
                doc.add_paragraph("_" * 60)

    doc.save(str(output_path))
    return output_path


# ==========================================
# PDF 匯出
# ==========================================

def _register_cjk_font():
    """註冊 CJK 字型供 reportlab 使用。

    優先使用 UnicodeCIDFont（STSong-Light），可同時正確渲染 ASCII 與中文字元。
    TTFont 方案（NotoSansCJK）雖字形更完整，但 ASCII 字元在此系統上有 null byte 問題，
    故降為備選。
    """
    from reportlab.pdfbase import pdfmetrics

    # 第一優先：reportlab 內建 CID 字型（ASCII + CJK 均正確渲染）
    try:
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        pass

    # 第二優先：系統 TTFont（部分環境 ASCII 字元可能顯示異常）
    from reportlab.pdfbase.ttfonts import TTFont
    font_paths = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
    ]
    for font_path in font_paths:
        if Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont("CJK", font_path, subfontIndex=4))
                return "CJK"
            except Exception:
                try:
                    pdfmetrics.registerFont(TTFont("CJK", font_path))
                    return "CJK"
                except Exception:
                    continue

    # 最後備選：無中文支援
    return "Helvetica"


def export_to_pdf(data: ReportData, output_path: Path) -> Path:
    """
    匯出稽核結果為 PDF 文件

    Args:
        data: 正規化的報告資料
        output_path: 輸出檔案路徑

    Returns:
        輸出檔案路徑
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    cjk_font = _register_cjk_font()

    # 建立文件
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
    )

    # 自訂樣式
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CJKTitle",
        parent=styles["Title"],
        fontName=cjk_font,
        fontSize=18,
        spaceAfter=20,
    )
    heading_style = ParagraphStyle(
        "CJKHeading",
        parent=styles["Heading1"],
        fontName=cjk_font,
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
    )
    heading2_style = ParagraphStyle(
        "CJKHeading2",
        parent=styles["Heading2"],
        fontName=cjk_font,
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "CJKBody",
        parent=styles["Normal"],
        fontName=cjk_font,
        fontSize=9,
        leading=14,
        spaceAfter=4,
    )

    elements = []

    # ---- 封面 ----
    elements.append(Spacer(1, 3 * cm))
    elements.append(Paragraph(data.title, title_style))
    elements.append(Spacer(1, 1 * cm))
    elements.append(Paragraph(f"產出時間: {data.generated_at}", body_style))
    elements.append(Paragraph(f"稽核規則數: {data.rules_audited}", body_style))
    elements.append(PageBreak())

    # ---- 統計摘要 ----
    elements.append(Paragraph("一、稽核統計摘要", heading_style))

    stats_data = [
        ["項目", "數量"],
        ["RISK_DETECTED (風險偵測)", str(data.summary_stats.get("RISK_DETECTED", 0))],
        ["COMPLIANT (合規)", str(data.summary_stats.get("COMPLIANT", 0))],
        ["UNCERTAIN (不確定)", str(data.summary_stats.get("UNCERTAIN", 0))],
        ["NOT_APPLICABLE (不適用)", str(data.summary_stats.get("NOT_APPLICABLE", 0))],
        ["HIGH 風險", str(data.severity_stats.get("HIGH", 0))],
        ["MEDIUM 風險", str(data.severity_stats.get("MEDIUM", 0))],
    ]

    stats_table = Table(stats_data, colWidths=[10 * cm, 4 * cm])
    stats_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), cjk_font),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#F2F2F2"), HexColor("#FFFFFF")]),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    elements.append(stats_table)
    elements.append(Spacer(1, 1 * cm))

    # ---- 摘要表格 ----
    if data.results:
        elements.append(Paragraph("二、逐規則摘要", heading_style))

        table_data = [["規則", "檢核項目", "判定結果", "風險等級", "條款參照"]]
        for result in data.results:
            rule = result.get("rule", {})
            compliance = result.get("compliance_analysis", {})
            table_data.append([
                rule.get("id", ""),
                rule.get("category", "")[:16],
                compliance.get("status", ""),
                compliance.get("severity", ""),
                compliance.get("clause_reference", "")[:20],
            ])

        summary_table = Table(table_data, colWidths=[2.5 * cm, 4 * cm, 3.5 * cm, 2.5 * cm, 4 * cm])
        table_style_cmds = [
            ("FONTNAME", (0, 0), (-1, -1), cjk_font),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#F2F2F2"), HexColor("#FFFFFF")]),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]

        # 狀態色彩
        for row_idx, result in enumerate(data.results, 1):
            status = result.get("compliance_analysis", {}).get("status", "")
            if status == "RISK_DETECTED":
                table_style_cmds.append(("TEXTCOLOR", (2, row_idx), (2, row_idx), HexColor("#CC0000")))
            elif status == "COMPLIANT":
                table_style_cmds.append(("TEXTCOLOR", (2, row_idx), (2, row_idx), HexColor("#008000")))

        summary_table.setStyle(TableStyle(table_style_cmds))
        elements.append(summary_table)
        elements.append(PageBreak())

        # ---- 逐規則詳細 ----
        elements.append(Paragraph("三、逐規則詳細報告", heading_style))

        for idx, result in enumerate(data.results, 1):
            rule = result.get("rule", {})
            query = result.get("query_generation", {})
            search = result.get("local_search_result", {})
            compliance = result.get("compliance_analysis", {})

            elements.append(Paragraph(
                f"{rule.get('id', '')} - {rule.get('category', '')}",
                heading2_style,
            ))

            # 規則資訊
            elements.append(Paragraph(f"<b>錯誤樣態:</b> {_escape_xml(rule.get('risk_pattern', ''))}", body_style))
            elements.append(Paragraph(f"<b>應採行動:</b> {_escape_xml(rule.get('action', ''))}", body_style))

            # 查詢
            if query:
                elements.append(Paragraph(
                    f"<b>查詢語句:</b> {_escape_xml(query.get('graph_query', ''))}",
                    body_style,
                ))

            # 合規分析
            status_text = compliance.get("status", "")
            status_color = "#CC0000" if status_text == "RISK_DETECTED" else "#008000" if status_text == "COMPLIANT" else "#333333"
            elements.append(Paragraph(
                f"<b>判定結果:</b> <font color='{status_color}'>{_escape_xml(status_text)}</font>",
                body_style,
            ))
            elements.append(Paragraph(f"<b>風險等級:</b> {_escape_xml(compliance.get('severity', ''))}", body_style))
            elements.append(Paragraph(f"<b>判斷理由:</b> {_escape_xml(compliance.get('reason', ''))}", body_style))

            evidence = compliance.get("evidence", "")
            if evidence:
                elements.append(Paragraph(f"<b>契約證據:</b> {_escape_xml(evidence[:500])}", body_style))

            elements.append(Paragraph(
                f"<b>建議行動:</b> {_escape_xml(compliance.get('recommendation', ''))}",
                body_style,
            ))

            elements.append(Spacer(1, 0.5 * cm))

    doc.build(elements)
    return output_path


def _escape_xml(text: str) -> str:
    """Escape XML 特殊字元（供 reportlab Paragraph 使用）"""
    if not text:
        return ""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


# ==========================================
# 高階匯出入口
# ==========================================

def export_report(
    source,
    output_dir: Path,
    formats: Optional[List[str]] = None,
    filename_prefix: str = "audit_report",
) -> List[Path]:
    """
    匯出稽核報告

    Args:
        source: 資料來源（目錄、JSON 檔、BatchAuditSummary）
        output_dir: 輸出目錄
        formats: 格式列表 ["docx", "pdf"]
        filename_prefix: 檔名前綴

    Returns:
        產生的檔案路徑列表
    """
    if formats is None:
        formats = ["docx"]

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    data = load_audit_results(source)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    generated_paths = []

    for fmt in formats:
        output_path = output_dir / f"{filename_prefix}_{timestamp}.{fmt}"

        if fmt == "docx":
            export_to_docx(data, output_path)
            generated_paths.append(output_path)
        elif fmt == "pdf":
            export_to_pdf(data, output_path)
            generated_paths.append(output_path)
        else:
            print(f"不支援的格式: {fmt}")

    return generated_paths


# ==========================================
# CLI
# ==========================================

def main():
    """報告匯出 CLI 入口"""
    parser = argparse.ArgumentParser(
        description="稽核報告匯出",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
範例:
  python report_export.py --input audit_reports/ --format docx
  python report_export.py --input audit_reports/batch_summary_xxx.json --format pdf
  python report_export.py --input audit_reports/ --format both
        """,
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="輸入來源: audit_reports/ 目錄或 batch_summary JSON 檔案",
    )
    parser.add_argument(
        "--format", "-f",
        choices=["pdf", "docx", "both"],
        default="docx",
        help="輸出格式 (預設: docx)",
    )
    parser.add_argument(
        "--output", "-o",
        default=".",
        help="輸出目錄 (預設: 當前目錄)",
    )
    parser.add_argument(
        "--prefix",
        default="audit_report",
        help="輸出檔名前綴 (預設: audit_report)",
    )

    args = parser.parse_args()

    formats = ["pdf", "docx"] if args.format == "both" else [args.format]

    print(f"載入稽核結果: {args.input}")

    try:
        paths = export_report(
            source=Path(args.input),
            output_dir=Path(args.output),
            formats=formats,
            filename_prefix=args.prefix,
        )

        for p in paths:
            print(f"報告已匯出: {p}")

    except Exception as e:
        print(f"匯出失敗: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
